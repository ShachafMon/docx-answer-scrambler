import docx
import re
import random
import logging
import sys
import copy
import subprocess
import tempfile
import shutil
import platform
from pathlib import Path
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import CONTENT_TYPE as CT
from docx.shared import Pt

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

if getattr(sys, "frozen", False):
    SCRIPT_DIR = Path(sys.executable).resolve().parent
else:
    SCRIPT_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = SCRIPT_DIR / "shuffled_output"

LETTERS = ["א", "ב", "ג", "ד", "ה","ו","ז"]
LETTER_PATTERN = re.compile(r'^\s*[\[]?\s*([א-ז])[\.\)]\)?')
NUM_PATTERN = re.compile(r'^\s*[\[]?\s*([1-5])[\.\)]\)?')
QUESTION_PATTERN = re.compile(r'^\s*שאלה\s+(\d+)\s*$')

INKSCAPE_DPI = 150


def setup_logger():
    logger = logging.getLogger("shuffle")
    logger.setLevel(logging.INFO)
    fmt = logging.Formatter("%(asctime)s | %(levelname)s | %(message)s", "%H:%M:%S")

    # Only warnings/errors interrupt the clean status UI (see ui_* below) —
    # everything else is narrated by ui_* prints instead of the logger.
    console = logging.StreamHandler(sys.stdout)
    console.setLevel(logging.WARNING)
    console.setFormatter(fmt)
    logger.addHandler(console)
    return logger


log = setup_logger()


def _enable_windows_ansi():
    try:
        import ctypes
        kernel32 = ctypes.windll.kernel32
        handle = kernel32.GetStdHandle(-11)
        mode = ctypes.c_uint32()
        if not kernel32.GetConsoleMode(handle, ctypes.byref(mode)):
            return False
        return bool(kernel32.SetConsoleMode(handle, mode.value | 0x0004))
    except Exception:
        return False


COLOR = sys.stdout.isatty() and (platform.system() != "Windows" or _enable_windows_ansi())

_CODES = {"green": "32", "red": "31", "cyan": "36", "yellow": "33", "bold": "1", "dim": "2"}


def style(text, *names):
    if not COLOR:
        return text
    codes = ";".join(_CODES[n] for n in names)
    return f"\033[{codes}m{text}\033[0m"


def ui_banner():
    width = 56
    print(style("=" * width, "cyan"))
    print(style("DOCX Answer Scrambler".center(width), "cyan", "bold"))
    print(style("=" * width, "cyan"))
    print()


def ui_file_start(idx, total, name):
    print(style(f"[{idx}/{total}]", "cyan", "bold") + f" {name}")


def ui_progress(n):
    sys.stdout.write(f"\r    shuffling... {n} question(s)")
    sys.stdout.flush()


def _ui_clear_line():
    sys.stdout.write("\r" + " " * 60 + "\r")


def ui_file_done(q_count, out_name):
    _ui_clear_line()
    print(style("    +", "green", "bold") + f" {q_count} question(s) shuffled -> {out_name}")


def ui_file_failed(err):
    _ui_clear_line()
    print(style("    x", "red", "bold") + f" failed: {err}")


def ui_summary(success, fail, out_dir):
    print()
    print(style("-" * 56, "dim"))
    line = f"  {success} succeeded"
    if fail:
        print(style(line + f", {fail} failed", "red", "bold"))
    else:
        print(style(line, "green", "bold"))
    print(f"  Output: {out_dir}")
    print(style("-" * 56, "dim"))


def get_para_text(p):
    return p.text.strip()


def label_match(text):
    return LETTER_PATTERN.match(text) or NUM_PATTERN.match(text)


def strip_label(text):
    m = label_match(text)
    return text[m.end():].strip() if m else text


def get_numid(paragraph):
    pPr = paragraph._p.pPr
    if pPr is None:
        return None
    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        return None
    numId_el = numPr.find(qn('w:numId'))
    if numId_el is None:
        return None
    return numId_el.get(qn('w:val'))


def has_math(paragraph):
    return paragraph._p.find(qn('m:oMath')) is not None


def rebuild_paragraph_text(paragraph, new_text):
    for run in list(paragraph.runs):
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def split_run_at(run, local_end):
    """Split run's text at local_end, keeping the first part in place and
    returning a new sibling run (same formatting) holding the remainder."""
    full_text = run.text
    run.text = full_text[:local_end]
    remainder = full_text[local_end:]
    new_r = copy.deepcopy(run._r)
    for t in new_r.findall(qn('w:t')):
        new_r.remove(t)
    new_t = OxmlElement('w:t')
    new_t.text = remainder
    new_t.set(qn('xml:space'), 'preserve')
    new_r.append(new_t)
    run._r.addnext(new_r)
    return new_r


def content_nodes_after_label(paragraph, match):
    """Non-pPr XML children that come after the label text, splitting the
    run in two if the label and answer body share the same run."""
    p = paragraph._p
    runs = paragraph.runs
    if not runs:
        return [c for c in p if c.tag != qn('w:pPr')]
    cum = 0
    content_start_el = None
    for r in runs:
        run_len = len(r.text)
        if cum + run_len > match.end():
            content_start_el = split_run_at(r, match.end() - cum)
            break
        cum += run_len
        if cum == match.end():
            content_start_el = r._r.getnext()
            break
    if content_start_el is None:
        content_start_el = runs[-1]._r.getnext()

    children = [c for c in p if c.tag != qn('w:pPr')]
    if content_start_el is None:
        return []
    pos = children.index(content_start_el)
    return children[pos:]


def extract_content_nodes(paragraph, match):
    """Detach and return the answer-body XML nodes for one option paragraph."""
    nodes = content_nodes_after_label(paragraph, match) if match else \
        [c for c in paragraph._p if c.tag != qn('w:pPr')]
    for n in nodes:
        n.getparent().remove(n)
    return nodes


def block_is_math(block_paras):
    return any(has_math(p) for p in block_paras)


def shuffle_block_nodes(block_paras, matches, rng):
    """Shuffle answer content by moving XML nodes, preserving formatting/equations."""
    extracted = [extract_content_nodes(p, m) for p, m in zip(block_paras, matches)]
    indices = list(range(len(extracted)))
    rng.shuffle(indices)
    for j, p in enumerate(block_paras):
        for node in extracted[indices[j]]:
            p._p.append(node)
    return indices


def shuffle_block_text(block_paras, rng):
    """Shuffle plain-text answer bodies by rewriting run text."""
    original_bodies = [strip_label(get_para_text(p)) for p in block_paras]
    indices = list(range(len(original_bodies)))
    rng.shuffle(indices)
    shuffled_bodies = [original_bodies[idx] for idx in indices]

    for j, p in enumerate(block_paras):
        orig_text = get_para_text(p)
        if LETTER_PATTERN.match(orig_text):
            label = LETTERS[j] if j < len(LETTERS) else str(j + 1)
            new_full_text = f"{label}. {shuffled_bodies[j]}"
        else:
            new_full_text = f"{j + 1}. {shuffled_bodies[j]}"
        rebuild_paragraph_text(p, new_full_text)
    return indices


def collect_label_block(paragraphs, i):
    block = []
    while i < len(paragraphs) and label_match(get_para_text(paragraphs[i])):
        block.append(paragraphs[i])
        i += 1
    return block, i


def collect_numid_block(paragraphs, i, numid):
    block = []
    while i < len(paragraphs) and get_numid(paragraphs[i]) == numid:
        block.append(paragraphs[i])
        i += 1
    return block, i


def is_genuine_answer_block(paragraphs, end_i):
    if end_i >= len(paragraphs):
        return True
    next_text = get_para_text(paragraphs[end_i])
    return next_text == "" or bool(QUESTION_PATTERN.search(next_text))


def convert_wmf_to_png_bytes(wmf_blob, dpi=INKSCAPE_DPI):
    """Convert a WMF blob to PNG bytes using Inkscape's independent WMF
    importer, which handles legacy metafile text-positioning far more
    reliably than LibreOffice's docx importer or ImageMagick/libwmf."""
    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        wmf_path = tmp_dir / "input.wmf"
        png_path = tmp_dir / "output.png"
        wmf_path.write_bytes(wmf_blob)

        result = subprocess.run(
            ["inkscape", str(wmf_path), "--export-type=png",
             f"--export-dpi={dpi}", "-o", str(png_path)],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode != 0 or not png_path.exists():
            raise RuntimeError(f"Inkscape failed to convert WMF: {result.stderr.strip()}")
        return png_path.read_bytes()


def replace_wmf_images_with_png(doc):
    """Find every WMF-typed image part in the docx, convert it to PNG via
    Inkscape, and swap the part's blob + content type in place so LibreOffice
    (or any viewer) renders the pre-rasterized PNG instead of parsing the WMF."""
    replaced = 0
    for rel_id, rel in list(doc.part.rels.items()):
        if "image" not in rel.reltype:
            continue
        target_ref = rel.target_ref.lower()
        if not target_ref.endswith(".wmf"):
            continue

        image_part = rel.target_part
        try:
            png_bytes = convert_wmf_to_png_bytes(image_part.blob)
        except Exception as e:
            log.warning(f"  Failed to convert WMF image ({target_ref}): {e}")
            continue

        image_part._blob = png_bytes
        image_part._content_type = CT.PNG
        if hasattr(image_part, "partname"):
            new_partname = image_part.partname.with_ext("png") \
                if hasattr(image_part.partname, "with_ext") else image_part.partname

        replaced += 1

    return replaced


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(OxmlElement('w:bidi'))

def add_answer_key_page(doc, answer_key):
    doc.add_page_break()

    heading = doc.add_paragraph()
    set_rtl(heading)
    heading_run = heading.add_run("מפתח תשובות")
    heading_run.bold = True
    heading_run.font.size = Pt(16)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table._tbl.tblPr.append(OxmlElement('w:bidiVisual'))
    header_cells = table.rows[0].cells
    header_cells[0].text = "שאלה"
    header_cells[1].text = "תשובה נכונה"
    for cell in header_cells:
        for p in cell.paragraphs:
            set_rtl(p)
            for run in p.runs:
                run.bold = True

    for entry in sorted(answer_key, key=lambda e: int(e["question"])):
        row_cells = table.add_row().cells
        row_cells[0].text = entry["question"]
        row_cells[1].text = entry["correct_letter"]
        for cell in row_cells:
            for p in cell.paragraphs:
                set_rtl(p)

def process_document(input_path: Path, output_path: Path, seed=None, progress_cb=None):
    rng = random.Random(seed)
    doc = docx.Document(str(input_path))

    replace_wmf_images_with_png(doc)

    paragraphs = doc.paragraphs

    answer_key = []
    current_question_num = None
    i = 0
    questions_found = 0

    while i < len(paragraphs):
        text = get_para_text(paragraphs[i])
        q_match = QUESTION_PATTERN.search(text)

        if q_match:
            if current_question_num:
                log.warning(f"  Question {current_question_num}: no answers detected to shuffle, skipped")
            current_question_num = q_match.group(1)
            i += 1
            continue

        if not current_question_num:
            i += 1
            continue

        lmatch = label_match(text)
        numid = get_numid(paragraphs[i])
        is_label_style = bool(lmatch)

        if is_label_style:
            block_paras, end_i = collect_label_block(paragraphs, i)
        elif numid is not None:
            block_paras, end_i = collect_numid_block(paragraphs, i, numid)
        else:
            i += 1
            continue

        if len(block_paras) < 2 or not is_genuine_answer_block(paragraphs, end_i):
            i = end_i
            continue

        i = end_i

        if is_label_style and not block_is_math(block_paras):
            indices = shuffle_block_text(block_paras, rng)
        elif is_label_style:
            matches = [label_match(get_para_text(p)) for p in block_paras]
            indices = shuffle_block_nodes(block_paras, matches, rng)
        else:
            indices = shuffle_block_nodes(block_paras, [None] * len(block_paras), rng)

        new_correct_pos = indices.index(0)
        answer_key.append({
            "question": current_question_num,
            "correct_letter": LETTERS[new_correct_pos] if new_correct_pos < len(LETTERS) else str(new_correct_pos + 1)
        })

        questions_found += 1
        if progress_cb:
            progress_cb(questions_found)
        current_question_num = None

    if current_question_num:
        log.warning(f"  Question {current_question_num}: no answers detected to shuffle, skipped")

    add_answer_key_page(doc, answer_key)

    doc.save(str(output_path))

    return questions_found


def convert_to_pdf(docx_path: Path, out_dir: Path):
    """Final PDF export via LibreOffice. Safe now because all WMF images were
    already rasterized to PNG by Inkscape before saving the docx, so
    LibreOffice never has to lay out WMF text itself."""
    result = subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
        capture_output=True, text=True, timeout=120
    )
    if result.returncode != 0:
        raise RuntimeError(result.stderr.strip() or "soffice failed")
    pdf_path = out_dir / f"{docx_path.stem}.pdf"
    if not pdf_path.exists():
        raise RuntimeError("soffice did not produce a PDF file")
    return pdf_path


INSTALL_INSTRUCTIONS = {
    "Linux": {
        "inkscape": "sudo apt install inkscape",
        "soffice": "sudo apt install libreoffice",
    },
    "Windows": {
        "inkscape": "winget install Inkscape.Inkscape   (or download from https://inkscape.org/release)",
        "soffice": "winget install TheDocumentFoundation.LibreOffice   (or download from https://www.libreoffice.org/download)",
    },
    "Darwin": {
        "inkscape": "brew install --cask inkscape   (or download from https://inkscape.org/release)",
        "soffice": "brew install --cask libreoffice   (or download from https://www.libreoffice.org/download)",
    },
}


def check_dependencies():
    missing = [tool for tool in ("inkscape", "soffice") if shutil.which(tool) is None]
    if missing:
        os_name = platform.system()
        instructions = INSTALL_INSTRUCTIONS.get(os_name, {})
        log.error(f"Missing required program(s): {', '.join(missing)}")
        for tool in missing:
            how = instructions.get(tool)
            if how:
                log.error(f"  Install {tool}: {how}")
        if os_name == "Windows":
            log.error("  After installing, make sure the install location was added to PATH "
                       "(reopen this window/terminal, or restart Windows, if it was just installed).")
        else:
            log.error("  After installing, make sure the commands are on PATH, then try again.")
        return False
    return True


def main():
    ui_banner()

    if not check_dependencies():
        return

    docx_files = [
        f for f in SCRIPT_DIR.glob("*.docx")
        if not f.name.startswith("~$")
    ]

    if not docx_files:
        print(style("No .docx files found in this folder!", "yellow", "bold"))
        return

    OUTPUT_DIR.mkdir(exist_ok=True)

    success_count = 0
    fail_count = 0

    for idx, file_path in enumerate(docx_files, start=1):
        ui_file_start(idx, len(docx_files), file_path.name)
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                tmp_docx_path = Path(tmp_dir) / f"{file_path.stem}_shuffled.docx"
                q_count = process_document(file_path, tmp_docx_path, seed=None,
                                            progress_cb=ui_progress)
                pdf_path = convert_to_pdf(tmp_docx_path, OUTPUT_DIR)

            ui_file_done(q_count, pdf_path.name)
            success_count += 1
        except Exception as e:
            ui_file_failed(e)
            fail_count += 1

    ui_summary(success_count, fail_count, OUTPUT_DIR)


if __name__ == "__main__":
    try:
        main()
    except Exception:
        logging.getLogger("shuffle").exception("Unexpected error")
    finally:
        if getattr(sys, "frozen", False):
            input("\nDone. Press Enter to close this window.")